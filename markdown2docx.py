#!/usr/bin/env python3
"""
Markdown to DOCX converter
Converts Markdown files to DOCX format with support for tables, headers, lists, and formatting.
"""

import os
import re
import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from bs4 import BeautifulSoup
import argparse


class Markdown2Docx:
    """
    Main converter class for converting Markdown to DOCX
    """
    
    def __init__(self, input_file, output_file=None):
        """
        Initialize the converter
        
        Args:
            input_file (str): Path to the input markdown file
            output_file (str, optional): Path for the output docx file. 
                                       If None, will use input filename with .docx extension
        """
        self.input_file = input_file
        self.output_file = output_file or self._get_default_output_file()
        self.document = Document()
        self._setup_document_styles()
        
    def _get_default_output_file(self):
        """Generate default output filename"""
        base_name = os.path.splitext(self.input_file)[0]
        return f"{base_name}.docx"
    
    def _setup_document_styles(self):
        """Setup document styles for consistent formatting"""
        styles = self.document.styles
        
        # Add custom styles if they don't exist
        try:
            styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        except:
            pass
            
    def _parse_markdown_to_html(self):
        """Convert markdown to HTML using markdown library"""
        with open(self.input_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # Configure markdown extensions
        md = markdown.Markdown(extensions=['tables', 'fenced_code', 'codehilite'])
        html_content = md.convert(markdown_content)
        
        return html_content
    
    def _convert_html_to_docx(self, html_content):
        """Convert HTML content to DOCX format"""
        soup = BeautifulSoup(html_content, 'html.parser')
        
        for element in soup.find_all(recursive=False):
            self._process_element(element)
    
    def _process_element(self, element):
        """Process HTML elements and add to DOCX document"""
        tag_name = element.name
        
        if tag_name == 'h1':
            self._add_heading(element.get_text(), 1)
        elif tag_name == 'h2':
            self._add_heading(element.get_text(), 2)
        elif tag_name == 'h3':
            self._add_heading(element.get_text(), 3)
        elif tag_name == 'h4':
            self._add_heading(element.get_text(), 4)
        elif tag_name == 'h5':
            self._add_heading(element.get_text(), 5)
        elif tag_name == 'h6':
            self._add_heading(element.get_text(), 6)
        elif tag_name == 'p':
            self._add_paragraph(element)
        elif tag_name == 'ul':
            self._add_list(element, is_ordered=False)
        elif tag_name == 'ol':
            self._add_list(element, is_ordered=True)
        elif tag_name == 'blockquote':
            self._add_blockquote(element)
        elif tag_name == 'table':
            self._add_table(element)
        elif tag_name == 'pre':
            self._add_code_block(element)
    
    def _add_heading(self, text, level):
        """Add heading to document"""
        paragraph = self.document.add_heading(text, level)
    
    def _add_paragraph(self, element):
        """Add paragraph with formatting"""
        text = element.get_text()
        if text.strip():
            paragraph = self.document.add_paragraph()
            self._process_inline_styles(paragraph, element)
    
    def _add_list(self, element, is_ordered=False):
        """Add list items to document"""
        for li in element.find_all('li', recursive=False):
            text = li.get_text().strip()
            if text:
                paragraph = self.document.add_paragraph(
                    text, 
                    style='List Number' if is_ordered else 'List Bullet'
                )
                self._process_inline_styles(paragraph, li)
    
    def _add_blockquote(self, element):
        """Add blockquote to document"""
        text = element.get_text().strip()
        if text:
            paragraph = self.document.add_paragraph(text)
            paragraph.style = 'Quote'
    
    def _add_table(self, element):
        """Add table to document"""
        rows = element.find_all('tr')
        if not rows:
            return
        
        # Get column count from first row
        first_row_cells = rows[0].find_all(['td', 'th'])
        col_count = len(first_row_cells)
        
        if col_count == 0:
            return
        
        # Create table
        table = self.document.add_table(rows=len(rows), cols=col_count)
        table.style = 'Table Grid'
        
        # Populate table
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for col_idx, cell in enumerate(cells[:col_count]):
                cell_text = cell.get_text().strip()
                table_cell = table.cell(row_idx, col_idx)
                table_cell.text = cell_text
    
    def _add_code_block(self, element):
        """Add code block to document"""
        code_element = element.find('code')
        if code_element:
            text = code_element.get_text()
        else:
            text = element.get_text()
        
        paragraph = self.document.add_paragraph(text)
        paragraph.style = 'Code'
        
        # Set monospace font
        if paragraph.runs:
            run = paragraph.runs[0]
            run.font.name = 'Courier New'
    
    def _process_inline_styles(self, paragraph, element):
        """Process inline styles like bold, italic, etc."""
        # This is a simplified version - could be enhanced
        text = element.get_text()
        paragraph.text = text
        
        # Basic formatting detection
        if '**' in text or '__' in text:
            # Handle bold text
            pass
        if '*' in text or '_' in text:
            # Handle italic text
            pass
    
    def eat_soup(self):
        """
        Main conversion method - processes the markdown file
        
        Returns:
            bool: True if conversion successful
        """
        try:
            # Parse markdown to HTML
            html_content = self._parse_markdown_to_html()
            
            # Convert HTML to DOCX
            self._convert_html_to_docx(html_content)
            
            return True
        except Exception as e:
            print(f"Error during conversion: {str(e)}")
            return False
    
    def save(self):
        """Save the DOCX document
        
        Returns:
            bool: True if save successful, False otherwise
        """
        try:
            self.document.save(self.output_file)
            print(f"Document saved successfully: {self.output_file}")
            return True
        except Exception as e:
            print(f"Error saving document: {str(e)}")
            return False
    
    def save_html_preview(self, html_file=None):
        """Save HTML preview of the processed markdown
        
        Args:
            html_file (str, optional): Path for HTML preview file
            
        Returns:
            bool: True if save successful, False otherwise
        """
        if not html_file:
            base_name = os.path.splitext(self.input_file)[0]
            html_file = f"{base_name}_preview.html"
        
        try:
            html_content = self._parse_markdown_to_html()
            with open(html_file, 'w', encoding='utf-8') as f:
                f.write(f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Preview: {os.path.basename(self.input_file)}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; }}
        table {{ border-collapse: collapse; margin: 20px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; }}
        th {{ background-color: #f2f2f2; }}
        pre {{ background-color: #f4f4f4; padding: 10px; overflow-x: auto; }}
        blockquote {{ border-left: 4px solid #ddd; padding-left: 20px; margin: 0; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>""")
            print(f"HTML preview saved: {html_file}")
            return True
        except Exception as e:
            print(f"Error saving HTML preview: {str(e)}")
            return False


def main():
    """CLI interface for the markdown2docx converter"""
    parser = argparse.ArgumentParser(description='Convert Markdown to DOCX')
    parser.add_argument('input', help='Input markdown file')
    parser.add_argument('-o', '--output', help='Output docx file (optional)')
    parser.add_argument('--html-preview', action='store_true', 
                       help='Also generate HTML preview')
    
    args = parser.parse_args()
    
    # Check if input file exists
    if not os.path.isfile(args.input):
        print(f"Error: Input file '{args.input}' not found")
        return
    
    # Create converter instance
    converter = Markdown2Docx(args.input, args.output)
    
    # Convert markdown to docx
    if converter.eat_soup():
        # Save the document
        if converter.save():
            print("Conversion completed successfully!")
        
        # Generate HTML preview if requested
        if args.html_preview:
            converter.save_html_preview()
    else:
        print("Conversion failed")


if __name__ == "__main__":
    main()